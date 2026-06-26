# Python script to generate a premium MS Word document for the statistics project
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    """Adds a heading with consistent sizing, bolding, and spacing."""
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(21, 101, 192) # Dark Blue theme (#1565C0)
    if level == 1:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    p.keep_with_next = True
    return p

# Main document setup
doc = Document()

# Set standard margins (1 inch / 2.54 cm all around)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Configure default paragraph format
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(33, 33, 33) # Charcoal for premium readability

# ----------------- 1. HEADER -----------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(18)
title_run = title_p.add_run("LAPORAN PROYEK DATA MINING (CLO-6)\nKLASIFIKASI NASABAH BANK UNTUK KAMPANYE PEMASARAN TERM DEPOSIT MENGGUNAKAN METODE REGRESI LOGISTIK BINOMIAL")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(21, 101, 192) # #1565C0

members_p = doc.add_paragraph()
members_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
members_p.paragraph_format.space_after = Pt(24)
members_run = members_p.add_run(
    "Dosen Pengampu: [Nama Dosen]\n\n"
    "Disusun Oleh (Kelompok 4):\n"
    "Muhammad Ilham Maulana (NIM: 120222XXXX)\n"
    "Dhana Zeta Pangertu (NIM: 120222XXXX)\n"
    "Muhammad Ade Sulistiansyah (NIM: 120222XXXX)\n"
    "Alfaayyadh Nezati Qasyim (NIM: 120222XXXX)\n\n"
    "PROGRAM STUDI S1 SISTEM INFORMASI\n"
    "FAKULTAS REKAYASA INDUSTRI, TELKOM UNIVERSITY\n"
    "2026"
)
members_run.font.size = Pt(11)
members_run.font.bold = True
members_run.font.color.rgb = RGBColor(66, 66, 66)

# Page Break after Title Page
doc.add_page_break()

# ----------------- 2. ABSTRAK & KATA KUNCI -----------------
abstrak_title = doc.add_paragraph()
abstrak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstrak_title.paragraph_format.space_after = Pt(6)
abstrak_title_run = abstrak_title.add_run("ABSTRAK")
abstrak_title_run.font.size = Pt(12)
abstrak_title_run.font.bold = True
abstrak_title_run.font.color.rgb = RGBColor(21, 101, 192)

abstrak_p = doc.add_paragraph()
abstrak_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstrak_p.paragraph_format.left_indent = Inches(0.4)
abstrak_p.paragraph_format.right_indent = Inches(0.4)
abstrak_p.paragraph_format.space_after = Pt(6)
abstrak_p.paragraph_format.line_spacing = 1.15
abstrak_run = abstrak_p.add_run(
    "Kampanye pemasaran merupakan salah satu instrumen penting bagi institusi perbankan untuk menarik nasabah baru guna melakukan pembukaan term deposit. Keberhasilan kampanye ini sangat dipengaruhi oleh penargetan nasabah yang tepat. Penelitian ini bertujuan untuk mengklasifikasikan nasabah bank yang berpotensi menyetujui pembukaan term deposit berdasarkan karakteristik sosio-demografis dan riwayat kontak kampanye pemasaran. Metode analisis yang digunakan adalah Regresi Logistik Binomial. Data yang digunakan berasal dari dataset Bank Marketing (bank.csv) yang terdiri dari 11.162 observasi dan 17 variabel. Evaluasi model dilakukan dengan membagi data menjadi 80% data latih dan 20% data uji. Hasil pemodelan menunjukkan nilai Akaike Information Criterion (AIC) sebesar 8078,8. Model regresi logistik berhasil memprediksi nasabah dengan tingkat akurasi sebesar 78,76%, presisi 79,87%, recall 73,16%, F1-score 76,37%, dan nilai Area Under Curve (AUC) sebesar 0,882. Variabel durasi panggilan (duration), status kontak sebelumnya (prev_contact), dan hasil kampanye sebelumnya (poutcome) menjadi prediktor paling signifikan dalam menentukan keputusan nasabah untuk mengambil term deposit. Implikasi praktis dari penelitian ini adalah bank dapat memprioritaskan penargetan pada nasabah yang memiliki durasi interaksi telepon yang lama serta nasabah yang memiliki respons positif pada kampanye sebelumnya untuk meningkatkan efektivitas pemasaran."
)
abstrak_run.font.italic = True
abstrak_run.font.size = Pt(10)

keyword_p = doc.add_paragraph()
keyword_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
keyword_p.paragraph_format.left_indent = Inches(0.4)
keyword_p.paragraph_format.right_indent = Inches(0.4)
keyword_p.paragraph_format.space_after = Pt(24)
kw_label = keyword_p.add_run("Kata Kunci: ")
kw_label.font.bold = True
kw_label.font.size = Pt(10)
kw_vals = keyword_p.add_run("Klasifikasi, Regresi Logistik, Bank Marketing, Term Deposit, RStudio")
kw_vals.font.italic = True
kw_vals.font.size = Pt(10)


# Abstract in English (for premium journal aesthetic)
abstract_title = doc.add_paragraph()
abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
abstract_title.paragraph_format.space_after = Pt(6)
abstract_title_run = abstract_title.add_run("ABSTRACT")
abstract_title_run.font.size = Pt(12)
abstract_title_run.font.bold = True
abstract_title_run.font.color.rgb = RGBColor(21, 101, 192)

abstract_p = doc.add_paragraph()
abstract_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abstract_p.paragraph_format.left_indent = Inches(0.4)
abstract_p.paragraph_format.right_indent = Inches(0.4)
abstract_p.paragraph_format.space_after = Pt(6)
abstract_p.paragraph_format.line_spacing = 1.15
abstract_run = abstract_p.add_run(
    "Marketing campaigns are a crucial tool for banking institutions to attract new clients to subscribe to term deposits. The success of these campaigns heavily relies on precise targeting. This study aims to classify bank customers who are likely to subscribe to term deposits based on their socio-demographic profiles and historical campaign interactions. The analytical method used is Binomial Logistic Regression. The dataset, derived from the Bank Marketing dataset (bank.csv), comprises 11,162 observations and 17 variables. Model evaluation was performed by splitting the dataset into 80% training and 20% testing sets. The logistic regression model achieved a training fit with an Akaike Information Criterion (AIC) of 8078.8. On the test set, the model yielded a classification accuracy of 78.76%, precision of 79.87%, recall of 73.16%, F1-score of 76.37%, and an Area Under Curve (AUC) of 0.882. Call duration (duration), previous contact status (prev_contact), and previous campaign outcome (poutcome) were identified as the most significant predictors of term deposit subscription. Practically, banks should prioritize customers with longer call durations and those who responded positively to previous campaigns to maximize marketing efficiency."
)
abstract_run.font.italic = True
abstract_run.font.size = Pt(10)

ekeyword_p = doc.add_paragraph()
ekeyword_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
ekeyword_p.paragraph_format.left_indent = Inches(0.4)
ekeyword_p.paragraph_format.right_indent = Inches(0.4)
ekeyword_p.paragraph_format.space_after = Pt(24)
ekw_label = ekeyword_p.add_run("Keywords: ")
ekw_label.font.bold = True
ekw_label.font.size = Pt(10)
ekw_vals = ekeyword_p.add_run("Classification, Logistic Regression, Bank Marketing, Term Deposit, RStudio")
ekw_vals.font.italic = True
ekw_vals.font.size = Pt(10)

# ----------------- 3. PENDAHULUAN -----------------
add_styled_heading(doc, "1. PENDAHULUAN", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Di tengah persaingan industri perbankan yang semakin kompetitif, penghimpunan dana pihak ketiga (DPK) menjadi salah satu prioritas utama untuk menjaga likuiditas bank. Produk deposito berjangka atau term deposit adalah salah satu instrumen paling efektif dalam mengumpulkan dana jangka panjang. Namun, memasarkan produk ini bukanlah hal yang mudah. Bank seringkali menghadapi tantangan dalam merancang strategi pemasaran langsung (direct marketing) yang efisien."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Kampanye pemasaran melalui telepon (telemarketing) merupakan metode yang umum digunakan untuk menawarkan term deposit secara langsung kepada nasabah individu. Namun, melakukan pemasaran secara serampangan atau massal tanpa segmentasi target yang jelas memiliki banyak kelemahan. Hal ini tidak hanya memerlukan biaya kampanye yang tinggi serta waktu staf yang banyak, tetapi juga memicu kejenuhan nasabah (customer fatigue) akibat menerima terlalu banyak panggilan telepon promosi yang tidak relevan dengan kebutuhan mereka."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Solusi dari masalah ini adalah pendekatan pemasaran berbasis data (data-driven marketing) yang memanfaatkan pemodelan prediktif untuk melakukan klasifikasi nasabah. Dengan memprediksi nasabah mana saja yang memiliki probabilitas tinggi untuk menerima tawaran term deposit berdasarkan karakteristik sosial-demografis (seperti usia, pekerjaan, status pernikahan, saldo) dan riwayat komunikasi (seperti durasi panggilan, kontak sebelumnya), bank dapat menyasar kelompok nasabah yang tepat. Pendekatan terarah ini terbukti mampu menghemat biaya operasional secara signifikan sekaligus meningkatkan rasio konversi kampanye pemasaran."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(12)
p.add_run(
    "Dalam laporan proyek ini, metode Regresi Logistik Binomial diterapkan pada dataset Bank Marketing (bank.csv). Model regresi logistik dipilih karena sangat andal untuk klasifikasi biner dan memberikan interpretabilitas parameter yang tinggi melalui nilai Odds Ratio. Penelitian ini bertujuan untuk mengidentifikasi variabel kunci yang mempengaruhi minat nasabah dan memetakan model evaluasi performa guna memberikan rekomendasi yang presisi bagi pengambil keputusan di bank."
)

# ----------------- 4. LANDASAN TEORI -----------------
add_styled_heading(doc, "2. LANDASAN TEORI", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Klasifikasi adalah salah satu tugas utama dalam data mining yang bertujuan untuk menugaskan data baru ke dalam kelas atau kategori yang telah ditentukan sebelumnya berdasarkan pola yang dipelajari dari data historis. Proyek ini berfokus pada masalah klasifikasi biner, di mana variabel target (Y) hanya memiliki dua kategori alternatif: sukses (menyatakan pembukaan term deposit, bernilai 1) atau tidak sukses (tidak membuka term deposit, bernilai 0)."
)

add_styled_heading(doc, "2.1 Regresi Logistik Binomial", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Regresi Logistik Binomial adalah metode analisis statistik yang digunakan untuk memodelkan hubungan antara satu variabel respon biner dengan satu atau lebih variabel prediktor (bisa kuantitatif maupun kualitatif). Tidak seperti regresi linier biasa yang berasumsi bahwa variabel terikat berdistribusi normal dan memiliki hubungan linear dengan variabel bebas, regresi logistik mengasumsikan variabel respon mengikuti distribusi Bernoulli."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Untuk memetakan nilai regresi linier (yang berkisar dari minus tak hingga sampai tak hingga) ke dalam rentang probabilitas [0, 1], regresi logistik menggunakan fungsi logistik (sigmoid). Model matematika untuk probabilitas kejadian sukses (Y=1) diberikan oleh formula berikut:"
)

# Formula block (centered, italicized)
f1 = doc.add_paragraph()
f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
f1.paragraph_format.space_before = Pt(6)
f1.paragraph_format.space_after = Pt(6)
f1_run = f1.add_run("P(Y = 1 | X) = e^(β₀ + β₁X₁ + ... + βₖXₖ) / [1 + e^(β₀ + β₁X₁ + ... + βₖXₖ)]       (Persamaan 1)")
f1_run.font.italic = True
f1_run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Untuk memudahkan interpretasi hubungan linear, persamaan di atas ditransformasikan ke dalam fungsi Logit (log dari odds), yang memodelkan logaritma natural dari perbandingan peluang sukses terhadap peluang gagal:"
)

f2 = doc.add_paragraph()
f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
f2.paragraph_format.space_before = Pt(6)
f2.paragraph_format.space_after = Pt(6)
f2_run = f2.add_run("ln[ P(Y=1) / (1 - P(Y=1)) ] = β₀ + β₁X₁ + ... + βₖXₖ       (Persamaan 2)")
f2_run.font.italic = True
f2_run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Di mana β₀ melambangkan intersep, dan βᵢ melambangkan koefisien regresi untuk variabel independen Xᵢ."
)

add_styled_heading(doc, "2.2 Odds Ratio (OR)", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Odds Ratio (OR) didefinisikan sebagai eksponensial dari koefisien regresi, yaitu OR = e^(βᵢ). Odds Ratio menunjukkan seberapa besar kecenderungan perubahan peluang terjadinya peristiwa target jika variabel prediktor Xᵢ naik sebesar satu satuan (dengan asumsi variabel lain bernilai konstan). Jika nilai OR > 1, maka variabel independen tersebut memperbesar peluang terjadinya deposit nasabah, sedangkan OR < 1 menunjukkan efek negatif yang mengurangi peluang terjadinya deposit."
)

add_styled_heading(doc, "2.3 Metrik Evaluasi Klasifikasi", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Kinerja model klasifikasi dievaluasi menggunakan Confusion Matrix, yaitu tabel kontingensi berukuran 2x2 yang membandingkan nilai aktual (Reference) dengan nilai prediksi (Prediction). Empat komponen utama tabel tersebut adalah True Positive (TP), True Negative (TN), False Positive (FP), dan False Negative (FN). Metrik evaluasi yang diturunkan meliputi:"
)

# Bullet list for metrics
bullets = [
    ("Akurasi: ", "Proporsi seluruh prediksi benar terhadap total data uji.\nFormula: Akurasi = (TP + TN) / (TP + TN + FP + FN)"),
    ("Presisi: ", "Proporsi ketepatan prediksi positif dibandingkan seluruh data yang diprediksi positif.\nFormula: Presisi = TP / (TP + FP)"),
    ("Recall (Sensitivitas): ", "Kemampuan model mendeteksi kelas positif secara benar dari seluruh nasabah yang aktualnya positif.\nFormula: Recall = TP / (TP + FN)"),
    ("F1-Score: ", "Nilai rata-rata harmonis yang menyeimbangkan presisi dan recall.\nFormula: F1 = 2 * (Presisi * Recall) / (Presisi + Recall)"),
    ("ROC dan AUC: ", "Kurva ROC menggambarkan trade-off antara True Positive Rate (Sensitivity) dan False Positive Rate (1-Specificity). Area Under Curve (AUC) mengukur kinerja diskriminasi model secara keseluruhan (0,5 = acak, 1,0 = sempurna).")
]
for title, desc in bullets:
    p_b = doc.add_paragraph(style='List Bullet')
    p_b.paragraph_format.space_after = Pt(4)
    run_t = p_b.add_run(title)
    run_t.font.bold = True
    p_b.add_run(desc)

# ----------------- 5. METODOLOGI -----------------
add_styled_heading(doc, "3. METODOLOGI", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Metodologi penelitian ini mencakup beberapa fase berurutan yang dirancang untuk menjamin validitas pemodelan klasifikasi. Tahapan tersebut meliputi:"
)

workflow_steps = [
    "Pengumpulan Data: Menggunakan data marketing perbankan dari file bank.csv.",
    "Analisis Data Deskriptif (EDA): Mengeksplorasi distribusi variabel dan hubungannya dengan variabel target menggunakan visualisasi boxplot dan barplot.",
    "Preprocessing Data: Mengubah variabel kategorikal menjadi numerik biner atau faktor berurutan, membuat variabel dummy baru, serta menyeleksi variabel.",
    "Pembagian Data: Membagi dataset secara acak menjadi 80% Data Latih (Training Set) dan 20% Data Uji (Test Set) menggunakan partition seed 42.",
    "Pelatihan Model: Menyusun model Regresi Logistik menggunakan fungsi glm() di R dengan basis distribusi binomial logit.",
    "Evaluasi & Interpretasi: Menguji model pada data uji menggunakan Confusion Matrix, ROC-AUC, serta melakukan analisis signifikansi koefisien dan Odds Ratio."
]
for idx, step in enumerate(workflow_steps, 1):
    p_w = doc.add_paragraph(style='List Bullet')
    p_w.paragraph_format.space_after = Pt(4)
    run_num = p_w.add_run(f"Langkah {idx}: ")
    run_num.font.bold = True
    p_w.add_run(step)

add_styled_heading(doc, "3.1 Preprocessing Detail", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Dataset asli terdiri dari 11.162 observasi dengan 17 variabel awal. Tahapan pembersihan dan transformasi data yang dilakukan di R meliputi:"
)

prep_details = [
    "deposit (variabel target) diubah dari kategorikal 'yes'/'no' menjadi biner 1/0.",
    "default, housing, dan loan ditransformasikan menjadi biner 1/0 berdasarkan keberadaan kredit.",
    "education diubah menjadi skala numerik bertingkat (primary = 1, secondary = 2, tertiary = 3, sisanya = 0).",
    "Variabel pdays (hari setelah kontak sebelumnya) yang bernilai -1 (artinya tidak pernah dihubungi) diubah menjadi variabel biner baru prev_contact (1 jika pdays != -1, artinya pernah dihubungi; 0 jika pdays = -1). Kolom pdays kemudian dihapus.",
    "month diubah menjadi nilai urutan bulan (1 s.d 12) dari 'jan' s.d 'dec'.",
    "job, marital, contact, dan poutcome diubah menjadi representasi numerik bertingkat menggunakan fungsi factor() lalu as.integer() di R."
]
for detail in prep_details:
    p_d = doc.add_paragraph(style='List Bullet')
    p_d.paragraph_format.space_after = Pt(4)
    p_d.add_run(detail)

# ----------------- 6. HASIL DAN DISKUSI -----------------
add_styled_heading(doc, "4. HASIL DAN DISKUSI", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Bagian ini memaparkan temuan dari analisis eksplorasi data (EDA) serta hasil pemodelan klasifikasi Regresi Logistik menggunakan program R. Seluruh pengujian statistik dilakukan pada tingkat signifikansi α = 5%."
)

add_styled_heading(doc, "4.1 Exploratory Data Analysis (EDA)", level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Berdasarkan visualisasi yang disajikan pada Gambar 1, Gambar 2, Gambar 3, dan Gambar 4, kita dapat mengidentifikasi pola awal nasabah:"
)

# Add EDA Grid Image
if os.path.exists("eda_grid.png"):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(4)
    img_p.add_run().add_picture("eda_grid.png", width=Inches(5.2))
    
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run("Gambar 1. Hasil EDA: Distribusi Target, Boxplot Usia, Saldo, dan Durasi vs Deposit")
    cap_run.font.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.bold = True

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_desc.paragraph_format.line_spacing = 1.15
p_desc.paragraph_format.space_after = Pt(6)
p_desc.add_run(
    "1. Distribusi Kelas Target: Dataset menunjukkan proporsi kelas yang sangat berimbang dengan 52,62% nasabah menolak deposit ('no') dan 47,38% menerima deposit ('yes'). Hal ini sangat ideal untuk pemodelan klasifikasi karena menghindarkan model dari masalah ketidakseimbangan kelas (class imbalance).\n"
    "2. Karakteristik Usia (Age): Usia nasabah berkisar dari 18 hingga 95 tahun dengan median sekitar 39 tahun. Terlihat pada boxplot bahwa distribusi usia nasabah yang menyetujui deposit memiliki rentang yang lebih lebar ke arah usia pensiun lanjut.\n"
    "3. Karakteristik Saldo (Balance): Sebagian besar nasabah memiliki saldo rata-rata di bawah 2.000 EUR. Terlihat bahwa nasabah yang melakukan deposit memiliki median saldo yang sedikit lebih tinggi dibandingkan dengan yang tidak.\n"
    "4. Karakteristik Durasi Panggilan (Duration): Variabel durasi menunjukkan perbedaan visual yang sangat kontras. Nasabah yang menyetujui deposit ('yes') memiliki median durasi panggilan telepon yang jauh lebih lama dibandingkan nasabah yang menolak ('no'). Durasi panggilan merupakan proksi dari minat komunikasi nasabah terhadap produk."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Selanjutnya, Gambar 2 (Job Proportion Plot) menyajikan proporsi nasabah yang sukses mengambil deposit berdasarkan kategori pekerjaan mereka:"
)

# Add Job Prop Image
if os.path.exists("job_prop.png"):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(4)
    img_p.add_run().add_picture("job_prop.png", width=Inches(4.5))
    
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run("Gambar 2. Proporsi Deposit Berdasarkan Jenis Pekerjaan Nasabah")
    cap_run.font.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Kategori pekerjaan dengan persentase kesuksesan pembukaan deposit tertinggi adalah Mahasiswa (student) dan Pensiunan (retired) dengan rasio keberhasilan melebihi 60%. Sementara itu, kelompok pekerja Blue-collar (pekerja kasar) dan Services (jasa) menunjukkan tingkat kesuksesan terendah, kemungkinan karena alokasi pendapatan mereka lebih diprioritaskan untuk pemenuhan kebutuhan primer daripada investasi jangka panjang."
)

add_styled_heading(doc, "4.2 Hasil Estimasi Model Regresi Logistik", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Setelah memisahkan data menjadi training (80%) dan testing (20%), model regresi logistik biner dilatih menggunakan R. Nilai AIC model adalah 8078,8. Ringkasan hasil estimasi koefisien, nilai p-value, dan Odds Ratio (OR) ditunjukkan pada Tabel 1 di bawah ini."
)

# Add Table 1 Caption
tab1_cap = doc.add_paragraph()
tab1_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
tab1_cap.paragraph_format.space_before = Pt(12)
tab1_cap.paragraph_format.space_after = Pt(4)
tab1_cap_run = tab1_cap.add_run("Tabel 1. Estimasi Parameter Model Regresi Logistik Binomial")
tab1_cap_run.font.bold = True
tab1_cap_run.font.size = Pt(10)
tab1_cap_run.font.color.rgb = RGBColor(21, 101, 192)

# Create Coefficients Table
# Headers: Variabel, Estimate, Std. Error, z value, Pr(>|z|), Odds Ratio (OR), Keterangan
coef_data = [
    ("Intercept", "-5.719", "0.371", "-15.398", "< 2e-16", "0.003", "Signifikan (***)"),
    ("age (Usia)", "0.008", "0.003", "3.073", "0.002", "1.008", "Signifikan (**)"),
    ("job (Pekerjaan)", "0.012", "0.009", "1.338", "0.181", "1.012", "Tidak Signifikan"),
    ("marital (Status Nikah)", "0.216", "0.050", "4.289", "1.80e-05", "1.241", "Signifikan (***)"),
    ("education (Pendidikan)", "0.180", "0.036", "4.974", "6.55e-07", "1.197", "Signifikan (***)"),
    ("default (Default Kredit)", "-0.488", "0.246", "-1.986", "0.047", "0.614", "Signifikan (*)"),
    ("balance (Saldo)", "2.94e-05", "9.19e-06", "3.198", "0.001", "1.000", "Signifikan (**)"),
    ("housing (Kredit Rumah)", "-0.905", "0.060", "-15.098", "< 2e-16", "0.405", "Signifikan (***)"),
    ("loan (Pinjaman Pribadi)", "-0.697", "0.089", "-7.804", "5.98e-15", "0.498", "Signifikan (***)"),
    ("contact (Metode Kontak)", "-0.570", "0.040", "-14.109", "< 2e-16", "0.566", "Signifikan (***)"),
    ("day (Hari Hubungi)", "-0.009", "0.003", "-2.762", "0.006", "0.991", "Signifikan (**)"),
    ("month (Bulan Hubungi)", "-0.018", "0.010", "-1.763", "0.078", "0.982", "Marginal (.)"),
    ("duration (Durasi Call)", "0.005", "1.29e-04", "38.806", "< 2e-16", "1.005", "Signifikan (***)"),
    ("campaign (Jml Kontak)", "-0.112", "0.014", "-7.945", "1.94e-15", "0.894", "Signifikan (***)"),
    ("previous (Kontak Sblm)", "-0.004", "0.015", "-0.235", "0.814", "0.996", "Tidak Signifikan"),
    ("poutcome (Hasil Sblm)", "1.028", "0.068", "15.112", "< 2e-16", "2.794", "Signifikan (***)"),
    ("prev_contact (Status Sblm)", "3.439", "0.181", "19.015", "< 2e-16", "31.160", "Signifikan (***)")
]

t1 = doc.add_table(rows=1 + len(coef_data), cols=7)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_margins(t1.rows[0].cells[0], top=120, bottom=120, left=180, right=180)

# Formatting Header Row
headers = ["Variabel", "Estimate", "Std. Error", "z value", "Pr(>|z|)", "Odds Ratio", "Keterangan"]
for i, name in enumerate(headers):
    cell = t1.cell(0, i)
    cell.text = name
    set_cell_background(cell, "1565C0")
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Populate Table Data
for row_idx, data in enumerate(coef_data, 1):
    for col_idx, text in enumerate(data):
        cell = t1.cell(row_idx, col_idx)
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        # Alternate background color
        if row_idx % 2 == 0:
            set_cell_background(cell, "F2F5FA")
        else:
            set_cell_background(cell, "FFFFFF")
        
        # Align columns
        if col_idx == 0 or col_idx == 6:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(6)

add_styled_heading(doc, "4.3 Interpretasi Koefisien & Analisis Odds Ratio", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Berdasarkan hasil pemodelan pada Tabel 1, berikut adalah analisis rinci pengaruh masing-masing variabel prediktor:"
)

# Bullet analysis of key findings
findings = [
    ("Status Pernikahan (marital) dan Pendidikan (education): ", "Memiliki koefisien positif yang signifikan. Nasabah dengan status pendidikan lebih tinggi (tertiary) dan status pernikahan stabil memiliki kecenderungan lebih tinggi untuk menyetujui deposit (OR masing-masing 1.197 dan 1.241)."),
    ("Beban Finansial (housing & loan): ", "Variabel housing (OR = 0.405) dan loan (OR = 0.498) memiliki efek negatif yang sangat signifikan. Nasabah dengan cicilan rumah memiliki peluang menyetujui deposit 59.5% lebih rendah, dan nasabah dengan pinjaman pribadi memiliki peluang 50.2% lebih rendah. Hal ini terjadi karena keterbatasan dana mengendap (disposable income) mereka."),
    ("Durasi Kontak (duration): ", "Sangat signifikan dengan nilai p-value < 2e-16. Setiap peningkatan durasi panggilan telepon sebesar 1 detik meningkatkan odds nasabah menyetujui pembukaan deposit sebesar 0.5% (OR = 1.005). Hubungan ini sangat logis karena panggilan telepon yang berlangsung lama umumnya mencerminkan interaksi interaktif di mana nasabah menunjukkan ketertarikan tinggi."),
    ("Jumlah Kontak Kampanye (campaign): ", "Memiliki koefisien negatif signifikan (OR = 0.894). Setiap kali nasabah dihubungi berulang kali dalam kampanye yang sama, peluang mereka untuk menyetujui deposit menurun sekitar 10.6%. Hal ini mengindikasikan bahwa taktik menghubungi nasabah secara agresif justru memicu penolakan."),
    ("Pengaruh Hubungan Sebelumnya (prev_contact & poutcome): ", "Variabel prev_contact memiliki pengaruh terbesar di dalam model dengan OR = 31,160. Nasabah yang pernah dihubungi dalam kampanye sebelumnya memiliki odds 31 kali lebih tinggi untuk membuka deposit dibanding nasabah baru. Terlebih lagi, jika respon kampanye sebelumnya sukses (poutcome, OR = 2,794), peluang keberhasilan pemasaran meningkat drastis.")
]
for title, desc in findings:
    p_f = doc.add_paragraph(style='List Bullet')
    p_f.paragraph_format.space_after = Pt(4)
    run_t = p_f.add_run(title)
    run_t.font.bold = True
    p_f.add_run(desc)

add_styled_heading(doc, "4.4 Evaluasi Performa Model", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Pengujian performa klasifikasi dilakukan pada 20% data uji (2.232 sampel). Confusion Matrix menggambarkan hasil prediksi riil model terhadap data aktual sebagai berikut:"
)

# Confusion Matrix Image
if os.path.exists("confusion_matrix.png"):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(4)
    img_p.add_run().add_picture("confusion_matrix.png", width=Inches(3.5))
    
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(12)
    cap_run = cap_p.add_run("Gambar 3. Visualisasi Confusion Matrix Hasil Prediksi Data Uji")
    cap_run.font.italic = True
    cap_run.font.size = Pt(9.5)
    cap_run.font.bold = True

# Confusion Matrix Table
tabcm_cap = doc.add_paragraph()
tabcm_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
tabcm_cap.paragraph_format.space_before = Pt(12)
tabcm_cap.paragraph_format.space_after = Pt(4)
tabcm_cap_run = tabcm_cap.add_run("Tabel 2. Confusion Matrix Regresi Logistik")
tabcm_cap_run.font.bold = True
tabcm_cap_run.font.size = Pt(10)
tabcm_cap_run.font.color.rgb = RGBColor(21, 101, 192)

t_cm = doc.add_table(rows=3, cols=3)
t_cm.alignment = WD_TABLE_ALIGNMENT.CENTER

t_cm.cell(0, 0).text = "Prediksi \\ Aktual"
t_cm.cell(0, 1).text = "Aktual Tidak (0)"
t_cm.cell(0, 2).text = "Aktual Ya (1)"
t_cm.cell(1, 0).text = "Prediksi Tidak (0)"
t_cm.cell(1, 1).text = "992 (True Negative / TN)"
t_cm.cell(1, 2).text = "281 (False Negative / FN)"
t_cm.cell(2, 0).text = "Prediksi Ya (1)"
t_cm.cell(2, 1).text = "193 (False Positive / FP)"
t_cm.cell(2, 2).text = "766 (True Positive / TP)"

# Format CM Table
for r_idx in range(3):
    for c_idx in range(3):
        cell = t_cm.cell(r_idx, c_idx)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9.5)
        if r_idx == 0 or c_idx == 0:
            set_cell_background(cell, "1565C0")
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            if r_idx == c_idx:
                set_cell_background(cell, "E3F2FD") # Light blue for correct predictions
                run.font.bold = True
            else:
                set_cell_background(cell, "FFEBEE") # Light red for wrong predictions

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Metrik evaluasi lengkap yang dihitung berdasarkan Confusion Matrix di atas disajikan pada Tabel 3 di bawah ini:"
)

# Table 3 Caption
tab3_cap = doc.add_paragraph()
tab3_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
tab3_cap.paragraph_format.space_before = Pt(12)
tab3_cap.paragraph_format.space_after = Pt(4)
tab3_cap_run = tab3_cap.add_run("Tabel 3. Kinerja Model Klasifikasi Data Uji")
tab3_cap_run.font.bold = True
tab3_cap_run.font.size = Pt(10)
tab3_cap_run.font.color.rgb = RGBColor(21, 101, 192)

# Create Evaluation Metrics Table
metrics_data = [
    ("Akurasi", "78.76%", "Proporsi prediksi yang tepat dari model secara keseluruhan."),
    ("Presisi", "79.87%", "Ketepatan prediksi nasabah berminat dari yang dipromosikan model."),
    ("Recall (Sensitivity)", "73.16%", "Kemampuan model menjaring nasabah yang berminat secara riil."),
    ("F1-Score", "76.37%", "Nilai seimbang antara presisi dan recall."),
    ("AUC (Area Under Curve)", "0.882", "Kemampuan pemisahan kelas oleh model (diskriminasi).")
]

t3 = doc.add_table(rows=1 + len(metrics_data), cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ["Metrik Evaluasi", "Nilai", "Definisi & Interpretasi"]
for i, name in enumerate(headers3):
    cell = t3.cell(0, i)
    cell.text = name
    set_cell_background(cell, "1565C0")
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_idx, data in enumerate(metrics_data, 1):
    for col_idx, text in enumerate(data):
        cell = t3.cell(row_idx, col_idx)
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if row_idx % 2 == 0:
            set_cell_background(cell, "F2F5FA")
        else:
            set_cell_background(cell, "FFFFFF")
        
        if col_idx == 1:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.bold = True
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(6)

# Add ROC Curve Image
if os.path.exists("roc_curve.png"):
    p_roc = doc.add_paragraph()
    p_roc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_roc.paragraph_format.space_before = Pt(6)
    p_roc.paragraph_format.space_after = Pt(4)
    p_roc.add_run().add_picture("roc_curve.png", width=Inches(3.8))
    
    cap_roc = doc.add_paragraph()
    cap_roc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_roc.paragraph_format.space_after = Pt(12)
    cap_roc_run = cap_roc.add_run("Gambar 4. Kurva Receiver Operating Characteristic (ROC) dan Nilai AUC")
    cap_roc_run.font.italic = True
    cap_roc_run.font.size = Pt(9.5)
    cap_roc_run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(12)
p.add_run(
    "Hasil akhir menunjukkan nilai AUC yang sangat tinggi yaitu 0,882. Secara umum, nilai AUC di atas 0,8 diklasifikasikan sebagai diskriminasi yang sangat baik (good classification). Hal ini menegaskan bahwa model regresi logistik binomial memiliki ketangguhan yang tinggi untuk diimplementasikan dalam sistem penargetan pemasaran di bank."
)

# ----------------- 7. SIMPULAN -----------------
add_styled_heading(doc, "5. SIMPULAN & REKOMENDASI", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Berdasarkan analisis data dan hasil pengujian model klasifikasi Regresi Logistik pada dataset Bank Marketing, dapat ditarik beberapa simpulan penting:"
)

conclusions = [
    "Model Regresi Logistik Binomial terbukti efektif untuk melakukan klasifikasi minat nasabah bank dengan tingkat Akurasi sebesar 78,76%, Presisi 79,87%, dan nilai AUC sebesar 0,882.",
    "Variabel durasi panggilan (duration) memiliki dampak positif yang konsisten terhadap kesuksesan pembukaan deposit. Setiap detik ekstra dari durasi panggilan telepon terbukti meningkatkan peluang konversi nasabah.",
    "Beban cicilan perumahan (housing) dan pinjaman pribadi (loan) secara signifikan mengurangi minat nasabah untuk menaruh uang mereka pada term deposit.",
    "Nasabah yang pernah dihubungi pada kampanye pemasaran sebelumnya, terutama yang merespon sukses, memiliki probabilitas konversi 31 kali lebih tinggi dibandingkan nasabah yang belum memiliki histori interaksi."
]
for idx, concl in enumerate(conclusions, 1):
    p_c = doc.add_paragraph(style='List Bullet')
    p_c.paragraph_format.space_after = Pt(4)
    run_num = p_c.add_run(f"Simpulan {idx}: ")
    run_num.font.bold = True
    p_c.add_run(concl)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Rekomendasi taktis untuk tim manajemen pemasaran bank adalah:"
)

recommendations = [
    "Prioritas Penargetan: Terapkan model scoring nasabah ini ke basis data untuk menyaring nasabah prospek tinggi sebelum melakukan telemarketing.",
    "Pemasaran Hubungan Pelanggan (CRM): Fokuskan upaya kampanye utama pada nasabah lama yang telah memiliki riwayat positif dengan bank, daripada terus-menerus menghubungi nasabah baru yang belum teruji.",
    "Optimalisasi Waktu Telemarketing: Latih staf telemarketing untuk mempertahankan percakapan yang ramah dan solutif demi memperpanjang durasi panggilan (duration), mengingat variabel ini merupakan indikator minat terkuat.",
    "Segmentasi Beban Kredit: Hindari pemborosan biaya pemasaran pada kelompok nasabah yang memiliki cicilan perumahan aktif atau pinjaman pribadi karena peluang konversi mereka sangat kecil."
]
for idx, rec in enumerate(recommendations, 1):
    p_r = doc.add_paragraph(style='List Bullet')
    p_r.paragraph_format.space_after = Pt(4)
    run_num = p_r.add_run(f"Rekomendasi {idx}: ")
    run_num.font.bold = True
    p_r.add_run(rec)

# ----------------- 8. REFERENSI -----------------
add_styled_heading(doc, "REFERENSI", level=1)

references = [
    "Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied Logistic Regression. Third Edition. John Wiley & Sons.",
    "Moro, S., Cortez, P., & Rita, P. (2014). A data-driven approach to predict the success of bank telemarketing. Decision Support Systems, 62, 22-31.",
    "Kuhn, M. (2008). Building Predictive Models in R Using the caret Package. Journal of Statistical Software, 28(5), 1-26.",
    "Robin, X., Turck, N., Hainard, A., et al. (2011). pROC: an open-source package for R and S+ to analyze and compare ROC curves. BMC Bioinformatics, 12, 77."
]
for ref in references:
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ref.paragraph_format.left_indent = Inches(0.4)
    p_ref.paragraph_format.first_line_indent = Inches(-0.4)
    p_ref.paragraph_format.space_after = Pt(4)
    p_ref.paragraph_format.line_spacing = 1.15
    p_ref.add_run(ref)

# ----------------- 9. LINK PROGRAM RSTUDIO -----------------
add_styled_heading(doc, "LINK PROGRAM RSTUDIO", level=1)

p_link = doc.add_paragraph()
p_link.paragraph_format.space_after = Pt(12)
p_link.add_run("Seluruh kode pemrosesan data, eksplorasi, pemodelan, dan visualisasi grafik di RStudio dapat diakses melalui tautan repositori GitHub berikut:\n")
run_url = p_link.add_run("https://github.com/ilhamulana/wawasan-global-tik-mg15")
run_url.font.underline = True
run_url.font.color.rgb = RGBColor(21, 101, 192)

# ----------------- 10. TABEL KONTRIBUSI ANGGOTA -----------------
add_styled_heading(doc, "TABEL KONTRIBUSI ANGGOTA", level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(6)
p.add_run(
    "Pembagian tugas dalam penyusunan laporan proyek dan pembuatan program klasifikasi regresi logistik ini dilakukan secara berimbang. Detail peran masing-masing anggota kelompok disajikan pada Tabel 4."
)

# Table 4 Caption
tab4_cap = doc.add_paragraph()
tab4_cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
tab4_cap.paragraph_format.space_before = Pt(12)
tab4_cap.paragraph_format.space_after = Pt(4)
tab4_cap_run = tab4_cap.add_run("Tabel 4. Rincian Kontribusi Anggota Kelompok 4")
tab4_cap_run.font.bold = True
tab4_cap_run.font.size = Pt(10)
tab4_cap_run.font.color.rgb = RGBColor(21, 101, 192)

# Headers: No, Nama Anggota, NIM, Kontribusi Deskriptif
contrib_data = [
    ("1", "Muhammad Ilham Maulana", "120222XXXX", "Melakukan pre-processing data (mengubah variabel kategorikal menjadi biner dan integer di R), melatih model regresi logistik menggunakan glm(), dan menulis bab Metodologi Laporan."),
    ("2", "Dhana Zeta Pangertu", "120222XXXX", "Mengunduh dataset bank.csv, melakukan analisis eksplorasi data (EDA) awal dengan visualisasi boxplot dan barplot di R, dan menulis bab Pendahuluan Laporan."),
    ("3", "Muhammad Ade Sulistiansyah", "120222XXXX", "Melakukan evaluasi performa model klasifikasi (menghitung nilai Confusion Matrix, presisi, recall, F1, serta menggambar Kurva ROC dan AUC), dan menulis bab Hasil dan Diskusi Laporan."),
    ("4", "Alfaayyadh Nezati Qasyim", "120222XXXX", "Melakukan pencarian referensi pustaka ilmiah, melakukan perapian format layout laporan (Word docx), serta menulis bab Landasan Teori dan Simpulan Laporan.")
]

t4 = doc.add_table(rows=1 + len(contrib_data), cols=4)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER

headers4 = ["No", "Nama Anggota", "NIM", "Deskripsi Kontribusi (Tugas Utama)"]
col_widths = [Inches(0.5), Inches(2.0), Inches(1.2), Inches(3.3)]

for i, name in enumerate(headers4):
    cell = t4.cell(0, i)
    cell.text = name
    set_cell_background(cell, "1565C0")
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Populate
for row_idx, data in enumerate(contrib_data, 1):
    for col_idx, text in enumerate(data):
        cell = t4.cell(row_idx, col_idx)
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        if row_idx % 2 == 0:
            set_cell_background(cell, "F2F5FA")
        else:
            set_cell_background(cell, "FFFFFF")
        
        if col_idx == 3:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
output_path = "Laporan_Klasifikasi_Nasabah_Bank.docx"
doc.save(output_path)
print(f"Document successfully created at {os.path.abspath(output_path)}")
